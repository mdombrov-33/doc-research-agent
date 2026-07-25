import argparse
from io import BytesIO
from pathlib import Path
from textwrap import wrap
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import pymupdf
from docx import Document as DocxDocument

from evals.schemas import DocumentSpec, FactLedger

DEFAULT_BENCHMARK_ROOT = Path(__file__).parent / "benchmark"


def render_corpus(ledger: FactLedger, output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    for document in ledger.documents:
        path = output_dir / document.filename
        suffix = path.suffix.lower()
        if suffix == ".txt":
            path.write_text(_plain_text(document), encoding="utf-8")
        elif suffix == ".docx":
            _write_docx(document, path)
        else:
            _write_pdf(document, path)


def _plain_text(document: DocumentSpec) -> str:
    metadata = [
        document.title,
        f"Published: {document.published_on.isoformat()}",
    ]
    if document.supersedes_document_id is not None:
        metadata.append(f"Supersedes: {document.supersedes_document_id}")

    sections = [f"{passage.heading}\n\n{passage.text}" for passage in document.passages]
    return "\n\n".join([*metadata, *sections]) + "\n"


def _write_docx(document: DocumentSpec, path: Path) -> None:
    artifact = DocxDocument()
    artifact.add_heading(document.title, level=0)
    artifact.add_paragraph(f"Published: {document.published_on.isoformat()}")
    if document.supersedes_document_id is not None:
        artifact.add_paragraph(f"Supersedes: {document.supersedes_document_id}")
    for passage in document.passages:
        artifact.add_heading(passage.heading, level=1)
        artifact.add_paragraph(passage.text)
    artifact.save(str(path))
    _normalize_docx(path)


def _write_pdf(document: DocumentSpec, path: Path) -> None:
    artifact = pymupdf.open()
    page = artifact.new_page()
    y = 72.0

    for line, size in _pdf_lines(document):
        if y > 760:
            page = artifact.new_page()
            y = 72.0
        page.insert_text((72, y), line, fontsize=size, fontname="helv")
        y += 22 if size >= 16 else 16

    artifact.set_metadata(
        {
            "title": document.title,
            "author": "Document Research Agent evaluation benchmark",
            "creationDate": f"D:{document.published_on.strftime('%Y%m%d')}000000Z",
            "modDate": f"D:{document.published_on.strftime('%Y%m%d')}000000Z",
        }
    )
    artifact.save(path, no_new_id=True)
    artifact.close()


def _pdf_lines(document: DocumentSpec) -> list[tuple[str, int]]:
    lines: list[tuple[str, int]] = [(document.title, 18)]
    lines.extend((line, 10) for line in wrap(f"Published: {document.published_on.isoformat()}", 88))
    if document.supersedes_document_id is not None:
        lines.extend(
            (line, 10) for line in wrap(f"Supersedes: {document.supersedes_document_id}", 88)
        )
    lines.append(("", 10))

    for passage in document.passages:
        lines.append((passage.heading, 14))
        lines.extend((line, 10) for line in wrap(passage.text, 88))
        lines.append(("", 10))
    return lines


def _normalize_docx(path: Path) -> None:
    source = BytesIO(path.read_bytes())
    normalized = BytesIO()
    with ZipFile(source) as archive, ZipFile(normalized, "w") as output:
        for original in sorted(archive.infolist(), key=lambda item: item.filename):
            item = ZipInfo(original.filename, date_time=(1980, 1, 1, 0, 0, 0))
            item.compress_type = ZIP_DEFLATED
            item.external_attr = original.external_attr
            item.create_system = original.create_system
            output.writestr(item, archive.read(original.filename))
    path.write_bytes(normalized.getvalue())


def main() -> int:
    parser = argparse.ArgumentParser(description="Render corpus files from the fact ledger")
    parser.add_argument("--root", type=Path, default=DEFAULT_BENCHMARK_ROOT)
    args = parser.parse_args()

    ledger_path = args.root / "fact_ledger.json"
    ledger = FactLedger.model_validate_json(ledger_path.read_text())
    render_corpus(ledger, args.root / "corpus")
    print(f"Rendered {len(ledger.documents)} corpus documents.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
