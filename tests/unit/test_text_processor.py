from types import SimpleNamespace

import fitz  # type: ignore[import-untyped]
import pytest
from docx import Document as DocxDocument

from src.core.document_processing import text_processor
from src.core.document_processing.text_processor import TextExtractor, extract_entities
from src.core.exceptions import UnsupportedFileTypeError


async def test_extract_txt(tmp_path):
    path = tmp_path / "note.txt"
    path.write_text("hello world", encoding="utf-8")
    text = await TextExtractor.extract_from_file(str(path), "note.txt")
    assert text == "hello world"


async def test_extract_docx(tmp_path):
    path = tmp_path / "doc.docx"
    doc = DocxDocument()
    doc.add_paragraph("first line")
    doc.add_paragraph("second line")
    doc.save(str(path))
    text = await TextExtractor.extract_from_file(str(path), "doc.docx")
    assert text == "first line\nsecond line"


async def test_extract_pdf(tmp_path):
    path = tmp_path / "doc.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "pdf body text")
    pdf.save(str(path))
    pdf.close()
    text = await TextExtractor.extract_from_file(str(path), "doc.pdf")
    assert "pdf body text" in text


async def test_unsupported_extension_rejected(tmp_path):
    path = tmp_path / "data.csv"
    path.write_text("a,b,c", encoding="utf-8")
    with pytest.raises(UnsupportedFileTypeError):
        await TextExtractor.extract_from_file(str(path), "data.csv")


async def test_extension_matched_case_insensitively(tmp_path):
    path = tmp_path / "NOTE.TXT"
    path.write_text("upper ext", encoding="utf-8")
    text = await TextExtractor.extract_from_file(str(path), "NOTE.TXT")
    assert text == "upper ext"


def test_extract_entities_dedupes(monkeypatch):
    # Two mentions of "Acme" must collapse to one entity.
    ents = [SimpleNamespace(text="Acme"), SimpleNamespace(text="Acme"), SimpleNamespace(text="Bob")]
    fake_nlp = lambda _text: SimpleNamespace(ents=ents)  # noqa: E731
    monkeypatch.setattr(text_processor, "get_spacy_model", lambda: fake_nlp)
    result = extract_entities("Acme hired Acme employee Bob")
    assert sorted(result) == ["Acme", "Bob"]
