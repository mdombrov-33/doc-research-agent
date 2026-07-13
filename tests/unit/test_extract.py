import fitz  # type: ignore[import-untyped]
import pytest
from docx import Document as DocxDocument

from src.core.exceptions import DocumentLimitError, UnsupportedFileTypeError
from src.core.ingestion.extract import extract_from_file

EXTRACTION_LIMITS = {"max_pdf_pages": 10, "max_extracted_characters": 100}


async def test_extract_txt(tmp_path):
    path = tmp_path / "note.txt"
    path.write_text("hello world", encoding="utf-8")
    text = await extract_from_file(str(path), "note.txt", **EXTRACTION_LIMITS)
    assert text == "hello world"


async def test_extract_docx(tmp_path):
    path = tmp_path / "doc.docx"
    doc = DocxDocument()
    doc.add_paragraph("first line")
    doc.add_paragraph("second line")
    doc.save(str(path))
    text = await extract_from_file(str(path), "doc.docx", **EXTRACTION_LIMITS)
    assert text == "first line\nsecond line"


async def test_extract_pdf(tmp_path):
    path = tmp_path / "doc.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "pdf body text")
    pdf.save(str(path))
    pdf.close()
    text = await extract_from_file(str(path), "doc.pdf", **EXTRACTION_LIMITS)
    assert "pdf body text" in text


async def test_unsupported_extension_rejected(tmp_path):
    path = tmp_path / "data.csv"
    path.write_text("a,b,c", encoding="utf-8")
    with pytest.raises(UnsupportedFileTypeError):
        await extract_from_file(str(path), "data.csv", **EXTRACTION_LIMITS)


async def test_extension_matched_case_insensitively(tmp_path):
    path = tmp_path / "NOTE.TXT"
    path.write_text("upper ext", encoding="utf-8")
    text = await extract_from_file(str(path), "NOTE.TXT", **EXTRACTION_LIMITS)
    assert text == "upper ext"


async def test_extract_pdf_rejects_too_many_pages(tmp_path):
    path = tmp_path / "doc.pdf"
    pdf = fitz.open()
    pdf.new_page()
    pdf.new_page()
    pdf.save(str(path))
    pdf.close()

    with pytest.raises(DocumentLimitError, match="PDF page limit"):
        await extract_from_file(
            str(path),
            "doc.pdf",
            max_pdf_pages=1,
            max_extracted_characters=100,
        )


async def test_extract_txt_rejects_too_many_characters(tmp_path):
    path = tmp_path / "note.txt"
    path.write_text("hello", encoding="utf-8")

    with pytest.raises(DocumentLimitError, match="Extracted text limit"):
        await extract_from_file(
            str(path),
            "note.txt",
            max_pdf_pages=10,
            max_extracted_characters=4,
        )
